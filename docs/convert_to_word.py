from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 讀取 markdown 內容
with open('docs/LXC_Docker_Deployment_Guide.md', 'r', encoding='utf-8') as f:
    content = f.read()

doc = Document()

# 標題
title = doc.add_heading('LXC Docker 完整部署指南', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副標題
doc.add_paragraph('專案名稱：Buddhist Education LINE Bot System')
doc.add_paragraph('版本：2026.01')
doc.add_paragraph('更新日期：2026-01-13')
doc.add_paragraph()

# 解析 markdown 並轉換
sections = content.split('## ')
for section in sections[1:]:
    lines = section.strip().split('\n')
    heading = lines[0].strip()
    
    doc.add_heading(heading, level=1)
    
    in_code_block = False
    code_content = []
    
    for line in lines[1:]:
        if line.startswith('```'):
            if in_code_block:
                # 輸出程式碼區塊
                code_para = doc.add_paragraph()
                code_para.style = 'No Spacing'
                for code_line in code_content:
                    run = code_para.add_run(code_line + '\n')
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
        elif in_code_block:
            code_content.append(line)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Quote'
        elif line.strip() and not line.startswith('---'):
            doc.add_paragraph(line)

doc.save('docs/LXC_Docker_部署指南.docx')
print('Word 文件已生成: docs/LXC_Docker_部署指南.docx')
