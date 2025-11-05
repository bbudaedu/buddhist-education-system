"""
Document Generator Module

This module handles the generation of Word and Excel documents for the new book summary system.
It creates formatted documents with book summaries and detailed information.
"""

import os
import logging
from datetime import datetime
from typing import List, Dict, Any, Tuple

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Warning: python-docx not installed. Word document generation will not work.")
    Document = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter
except ImportError:
    print("Warning: openpyxl not installed. Excel document generation will not work.")
    Workbook = None


class DocumentGenerator:
    """
    Handles generation of Word and Excel documents for book summaries.
    
    This class creates formatted documents containing book information and summaries,
    following the requirements for the new book summary email system.
    """
    
    def __init__(self, logger: logging.Logger = None):
        """
        Initialize the DocumentGenerator.
        
        Args:
            logger: Logger instance for logging operations
        """
        self.logger = logger or logging.getLogger(__name__)
        self.word_doc = None
        self.excel_workbook = None
        self.excel_worksheet = None
        self.book_count = 0
        
        # Check if required libraries are available
        if Document is None:
            self.logger.warning("python-docx not available. Word document generation disabled.")
        if Workbook is None:
            self.logger.warning("openpyxl not available. Excel document generation disabled.")
    
    def create_word_document(self) -> Document:
        """
        Create a new Word document with title and date.
        
        Returns:
            Document: The created Word document
            
        Raises:
            ImportError: If python-docx is not installed
        """
        if Document is None:
            raise ImportError("python-docx is required for Word document generation")
        
        self.logger.info("Creating new Word document")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('新書簡介', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%Y年%m月%d日')
        date_para = doc.add_paragraph(current_date)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator line
        doc.add_paragraph('=' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add empty line
        doc.add_paragraph()
        
        self.word_doc = doc
        self.logger.info("Word document created successfully")
        return doc
    
    def create_excel_document(self) -> Workbook:
        """
        Create a new Excel workbook with headers.
        
        Returns:
            Workbook: The created Excel workbook
            
        Raises:
            ImportError: If openpyxl is not installed
        """
        if Workbook is None:
            raise ImportError("openpyxl is required for Excel document generation")
        
        self.logger.info("Creating new Excel workbook")
        
        wb = Workbook()
        ws = wb.active
        ws.title = "新書詳細資料"
        
        # Define headers
        headers = [
            '書號', '書名', '作者', 'PDF檔名', '檔案大小(MB)', 
            '處理方式', '摘要', '下載連結', '處理時間'
        ]
        
        # Add headers to worksheet
        ws.append(headers)
        
        # Format header row
        header_font = Font(bold=True, size=12, color='FFFFFF')
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Set column widths
        column_widths = {
            'A': 12,  # 書號
            'B': 30,  # 書名
            'C': 25,  # 作者
            'D': 25,  # PDF檔名
            'D': 15,  # 檔案大小
            'E': 12,  # 處理方式
            'F': 50,  # 摘要
            'G': 40,  # 下載連結
            'H': 20   # 處理時間
        }
        
        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width
        
        # Set row height for header
        ws.row_dimensions[1].height = 25
        
        self.excel_workbook = wb
        self.excel_worksheet = ws
        self.logger.info("Excel workbook created successfully")
        return wb
    
    def add_book_to_word(self, title: str, summary: str) -> None:
        """
        Add a book entry to the Word document.
        
        Args:
            title: Book title
            summary: Book summary text
        """
        if self.word_doc is None:
            self.logger.error("Word document not initialized. Call create_word_document() first.")
            return
        
        self.logger.debug(f"Adding book to Word document: {title}")
        
        # Clean summary - remove AI-generated prefixes
        cleaned_summary = self._clean_summary(summary)
        
        # Add book title (bold, 14pt)
        title_para = self.word_doc.add_paragraph()
        title_run = title_para.add_run(f'【{title}】')
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        # Add summary
        summary_para = self.word_doc.add_paragraph(cleaned_summary)
        summary_para.style = 'Normal'
        
        # Add blank line for spacing
        self.word_doc.add_paragraph()
        
        self.book_count += 1
        self.logger.debug(f"Book added to Word document. Total books: {self.book_count}")
    
    def _clean_summary(self, summary: str) -> str:
        """
        Clean summary text by removing AI-generated prefixes and unwanted text.
        
        Args:
            summary: Original summary text
            
        Returns:
            Cleaned summary text
        """
        if not summary:
            return summary
        
        # Remove common AI-generated prefixes
        prefixes_to_remove = [
            "好的，這是一份根據您提供的書籍內容生成的 300 字摘要：",
            "好的，這是一份根據您提供的書籍內容生成的300字摘要：",
            "好的，這是一份根據您提供的書籍內容生成的摘要：",
            "根據您提供的書籍內容，以下是300字摘要：",
            "以下是根據書籍內容生成的300字摘要：",
            "這是一份300字的書籍摘要：",
            "書籍摘要如下：",
            "摘要：",
        ]
        
        cleaned = summary.strip()
        
        for prefix in prefixes_to_remove:
            if cleaned.startswith(prefix):
                cleaned = cleaned[len(prefix):].strip()
                break
        
        return cleaned
    
    def _extract_book_code(self, filename: str, title: str) -> str:
        """
        Extract book code from filename or title.
        
        Args:
            filename: PDF filename
            title: Book title
            
        Returns:
            Book code (e.g., CH113-01)
        """
        if not filename:
            return title
        
        # Remove file extension
        name_without_ext = os.path.splitext(filename)[0]
        
        # Try to extract format: CH826-21-01-001 -> CH826-21
        parts = name_without_ext.split('-')
        if len(parts) >= 2:
            return f"{parts[0]}-{parts[1]}"
        
        return name_without_ext
    
    def _extract_author_info(self, title: str, book_data: Dict[str, Any]) -> str:
        """
        Extract author information from title or book data.
        
        Args:
            title: Book title
            book_data: Book data dictionary
            
        Returns:
            Author information
        """
        # Check if book_data has author field
        if 'author' in book_data:
            return book_data['author']
        
        # Try to extract from title - look for common patterns
        # Example: "全部佛法的綱要 道源長老 講述＼施旺坤 敬記"
        if '講述' in title or '敬記' in title or '編著' in title:
            # Split by common separators and try to find author part
            parts = title.split()
            for i, part in enumerate(parts):
                if any(keyword in part for keyword in ['講述', '敬記', '編著', '著']):
                    # Take the part before the keyword as potential author
                    if i > 0:
                        return parts[i-1] + ' ' + part
                    else:
                        return part
        
        # Default fallback
        return "未知作者"
    
    def save_word_document(self, output_dir: str = ".") -> str:
        """
        Save the Word document with the specified filename format and enhanced file system error handling.
        
        Args:
            output_dir: Directory to save the document
            
        Returns:
            str: Path to the saved document
            
        Raises:
            ValueError: If Word document is not initialized
            PermissionError: If no write permission to output directory
            OSError: If file system error occurs
        """
        if self.word_doc is None:
            raise ValueError("Word document not initialized. Call create_word_document() first.")
        
        # Generate filename with current date
        current_date = datetime.now().strftime('%Y-%m-%d')
        filename = f"新書簡介_{current_date}.docx"
        filepath = os.path.join(output_dir, filename)
        
        try:
            # Ensure output directory exists and is writable
            if not os.path.exists(output_dir):
                try:
                    os.makedirs(output_dir, exist_ok=True)
                    self.logger.info(f"Created output directory: {output_dir}")
                except PermissionError as e:
                    raise PermissionError(f"No permission to create directory: {output_dir} - {e}")
                except OSError as e:
                    raise OSError(f"Failed to create directory: {output_dir} - {e}")
            
            # Check if directory is writable
            if not os.access(output_dir, os.W_OK):
                raise PermissionError(f"No write permission to directory: {output_dir}")
            
            # Check if file already exists and handle it
            if os.path.exists(filepath):
                self.logger.warning(f"File already exists, will overwrite: {filepath}")
                # Check if existing file is writable
                if not os.access(filepath, os.W_OK):
                    raise PermissionError(f"No write permission to existing file: {filepath}")
            
            # Check available disk space (basic check)
            try:
                stat = os.statvfs(output_dir) if hasattr(os, 'statvfs') else None
                if stat:
                    available_space = stat.f_bavail * stat.f_frsize
                    if available_space < 10 * 1024 * 1024:  # Less than 10MB
                        self.logger.warning(f"Low disk space: {available_space / (1024*1024):.1f} MB available")
            except:
                pass  # Ignore disk space check errors on Windows
            
            # Save the document
            self.word_doc.save(filepath)
            
            # Verify the file was created successfully
            if not os.path.exists(filepath):
                raise OSError(f"File was not created successfully: {filepath}")
            
            file_size = os.path.getsize(filepath)
            if file_size == 0:
                raise OSError(f"Created file is empty: {filepath}")
            
            self.logger.info(f"Word document saved successfully: {filepath} ({file_size} bytes)")
            return filepath
            
        except (PermissionError, OSError) as fs_error:
            self.logger.error(f"File system error saving Word document: {fs_error}")
            raise fs_error
        except Exception as e:
            self.logger.error(f"Unexpected error saving Word document: {e}")
            raise OSError(f"Failed to save Word document: {e}")
    
    def generate_word_document(self, books_data: List[Dict[str, Any]], output_dir: str = ".") -> str:
        """
        Generate a complete Word document with all book summaries.
        
        Args:
            books_data: List of dictionaries containing book information
                       Each dict should have 'title' and 'summary' keys
            output_dir: Directory to save the document
            
        Returns:
            str: Path to the saved document
        """
        self.logger.info(f"Generating Word document with {len(books_data)} books")
        
        # Create new document
        self.create_word_document()
        
        # Add each book
        for book_data in books_data:
            title = book_data.get('title', 'Unknown Title')
            summary = book_data.get('summary', 'No summary available')
            self.add_book_to_word(title, summary)
        
        # Save document
        filepath = self.save_word_document(output_dir)
        
        self.logger.info(f"Word document generation completed: {filepath}")
        return filepath
    
    def add_book_to_excel(self, book_data: Dict[str, Any], index: int) -> None:
        """
        Add a book data row to the Excel worksheet.
        
        Args:
            book_data: Dictionary containing book information
            index: Sequential number for the book
        """
        if self.excel_worksheet is None:
            self.logger.error("Excel worksheet not initialized. Call create_excel_document() first.")
            return
        
        self.logger.debug(f"Adding book to Excel: {book_data.get('title', 'Unknown')}")
        
        # Extract data with defaults
        title = book_data.get('title', 'Unknown Title')
        filename = book_data.get('filename', book_data.get('pdf_filename', 'Unknown'))
        file_size_bytes = book_data.get('file_size_bytes', 0)
        file_size_mb = round(file_size_bytes / (1024 * 1024), 2) if file_size_bytes > 0 else 0
        processing_method = book_data.get('processing_method', 'Unknown')
        summary = book_data.get('summary', 'No summary available')
        download_url = book_data.get('pdf_url', book_data.get('download_url', ''))
        timestamp = book_data.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        
        # Extract book code from title (e.g., CH113-01 from filename CH113-01-01-001.pdf)
        book_code = self._extract_book_code(filename, title)
        
        # Extract author information from title or other sources
        author = self._extract_author_info(title, book_data)
        
        # Clean summary
        cleaned_summary = self._clean_summary(summary)
        
        # Convert processing method to Chinese
        method_mapping = {
            'pdf_extract': 'PDF提取',
            'pdf_extraction': 'PDF提取',
            'google_search': 'Google搜尋',
            'pdf': 'PDF提取',
            'search': 'Google搜尋'
        }
        processing_method_cn = method_mapping.get(processing_method.lower(), processing_method)
        
        # Prepare row data
        row_data = [
            book_code,                # 書號
            title,                    # 書名
            author,                   # 作者
            filename,                 # PDF檔名
            file_size_mb,            # 檔案大小(MB)
            processing_method_cn,     # 處理方式
            cleaned_summary,          # 摘要
            download_url,             # 下載連結
            timestamp                 # 處理時間
        ]
        
        # Add row to worksheet
        self.excel_worksheet.append(row_data)
        
        # Format the newly added row
        row_num = self.excel_worksheet.max_row
        
        # Set alignment for all cells in the row
        for col_num in range(1, len(row_data) + 1):
            cell = self.excel_worksheet.cell(row=row_num, column=col_num)
            
            if col_num == 1:  # 書號 - center aligned
                cell.alignment = Alignment(horizontal='center', vertical='center')
            elif col_num == 5:  # 檔案大小 - right aligned
                cell.alignment = Alignment(horizontal='right', vertical='center')
            elif col_num in [6, 7]:  # 摘要, 下載連結 - wrap text
                cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            else:  # Other columns - left aligned
                cell.alignment = Alignment(horizontal='left', vertical='center')
        
        # Set row height for better readability
        self.excel_worksheet.row_dimensions[row_num].height = 30
        
        self.logger.debug(f"Book added to Excel. Row: {row_num}")
    
    def save_excel_document(self, output_dir: str = ".") -> str:
        """
        Save the Excel workbook with the specified filename format and enhanced file system error handling.
        
        Args:
            output_dir: Directory to save the document
            
        Returns:
            str: Path to the saved document
            
        Raises:
            ValueError: If Excel workbook is not initialized
            PermissionError: If no write permission to output directory
            OSError: If file system error occurs
        """
        if self.excel_workbook is None:
            raise ValueError("Excel workbook not initialized. Call create_excel_document() first.")
        
        # Generate filename with current date
        current_date = datetime.now().strftime('%Y-%m-%d')
        filename = f"新書詳細資料_{current_date}.xlsx"
        filepath = os.path.join(output_dir, filename)
        
        try:
            # Ensure output directory exists and is writable
            if not os.path.exists(output_dir):
                try:
                    os.makedirs(output_dir, exist_ok=True)
                    self.logger.info(f"Created output directory: {output_dir}")
                except PermissionError as e:
                    raise PermissionError(f"No permission to create directory: {output_dir} - {e}")
                except OSError as e:
                    raise OSError(f"Failed to create directory: {output_dir} - {e}")
            
            # Check if directory is writable
            if not os.access(output_dir, os.W_OK):
                raise PermissionError(f"No write permission to directory: {output_dir}")
            
            # Check if file already exists and handle it
            if os.path.exists(filepath):
                self.logger.warning(f"File already exists, will overwrite: {filepath}")
                # Check if existing file is writable
                if not os.access(filepath, os.W_OK):
                    raise PermissionError(f"No write permission to existing file: {filepath}")
            
            # Check available disk space (basic check)
            try:
                stat = os.statvfs(output_dir) if hasattr(os, 'statvfs') else None
                if stat:
                    available_space = stat.f_bavail * stat.f_frsize
                    if available_space < 10 * 1024 * 1024:  # Less than 10MB
                        self.logger.warning(f"Low disk space: {available_space / (1024*1024):.1f} MB available")
            except:
                pass  # Ignore disk space check errors on Windows
            
            # Save the workbook
            self.excel_workbook.save(filepath)
            
            # Verify the file was created successfully
            if not os.path.exists(filepath):
                raise OSError(f"File was not created successfully: {filepath}")
            
            file_size = os.path.getsize(filepath)
            if file_size == 0:
                raise OSError(f"Created file is empty: {filepath}")
            
            self.logger.info(f"Excel document saved successfully: {filepath} ({file_size} bytes)")
            return filepath
            
        except (PermissionError, OSError) as fs_error:
            self.logger.error(f"File system error saving Excel document: {fs_error}")
            raise fs_error
        except Exception as e:
            self.logger.error(f"Unexpected error saving Excel document: {e}")
            raise OSError(f"Failed to save Excel document: {e}")
    
    def generate_excel_document(self, books_data: List[Dict[str, Any]], output_dir: str = ".") -> str:
        """
        Generate a complete Excel document with all book data.
        
        Args:
            books_data: List of dictionaries containing book information
            output_dir: Directory to save the document
            
        Returns:
            str: Path to the saved document
        """
        self.logger.info(f"Generating Excel document with {len(books_data)} books")
        
        # Create new workbook
        self.create_excel_document()
        
        # Add each book
        for index, book_data in enumerate(books_data, 1):
            self.add_book_to_excel(book_data, index)
        
        # Save document
        filepath = self.save_excel_document(output_dir)
        
        self.logger.info(f"Excel document generation completed: {filepath}")
        return filepath
    
    def generate_both_documents(self, books_data: List[Dict[str, Any]], output_dir: str = ".") -> Tuple[str, str]:
        """
        Generate both Word and Excel documents.
        
        Args:
            books_data: List of dictionaries containing book information
            output_dir: Directory to save the documents
            
        Returns:
            Tuple[str, str]: Paths to the saved Word and Excel documents
        """
        self.logger.info(f"Generating both documents with {len(books_data)} books")
        
        word_path = self.generate_word_document(books_data, output_dir)
        excel_path = self.generate_excel_document(books_data, output_dir)
        
        self.logger.info("Both documents generated successfully")
        return word_path, excel_path


# Example usage and testing
if __name__ == "__main__":
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Sample book data for testing
    sample_books = [
        {
            'title': '佛教基礎教義',
            'summary': '本書深入淺出地介紹了佛教的基本教義，包括四聖諦、八正道等核心概念。作者以現代語言詮釋古老智慧，適合初學者閱讀。書中結合實際案例，幫助讀者理解佛法在日常生活中的應用，是一本很好的佛教入門書籍。',
            'filename': 'CH826-21-01-001.pdf',
            'file_size_bytes': 2048576,  # 2MB
            'processing_method': 'pdf_extract',
            'pdf_url': 'https://example.com/CH826-21-01-001.pdf',
            'timestamp': '2024-01-15 10:30:00'
        },
        {
            'title': '禪修指導手冊',
            'summary': '這是一本實用的禪修指導書，詳細說明了各種禪修方法和技巧。從基礎的呼吸觀察到高深的內觀修行，循序漸進地引導讀者進入禪修的世界。書中包含豐富的實修經驗分享，對於想要深入學習禪修的人來說是不可多得的參考資料。',
            'filename': 'CH827-22-01-001.pdf',
            'file_size_bytes': 35651584,  # 34MB
            'processing_method': 'google_search',
            'pdf_url': 'https://example.com/CH827-22-01-001.pdf',
            'timestamp': '2024-01-15 11:45:00'
        }
    ]
    
    # Test document generation
    try:
        generator = DocumentGenerator()
        
        # Generate both documents
        word_path, excel_path = generator.generate_both_documents(sample_books, "test_output")
        
        print(f"Word document generated: {word_path}")
        print(f"Excel document generated: {excel_path}")
        
    except ImportError as e:
        print(f"Missing required library: {e}")
        print("Please install required packages: pip install python-docx openpyxl")
    except Exception as e:
        print(f"Error during document generation: {e}")