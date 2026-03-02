#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Generator Module for Website Monitoring
文件生成模組

This module handles Excel file generation for website monitoring content.
Extends existing document generation functionality for monitoring data.
"""

import os
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


class DocumentGenerator:
    """
    Document generator for creating Excel files from monitoring data
    
    Handles:
    - Excel file creation for different content types
    - Multi-sheet Excel files
    - Data formatting and styling
    - File management and organization
    """
    
    def __init__(self, output_dir: str = "generated_documents", logger: Optional[logging.Logger] = None):
        """
        Initialize DocumentGenerator
        
        Args:
            output_dir: Directory for generated documents
            logger: Logger instance for logging operations
        """
        self.logger = logger or logging.getLogger(__name__)
        self.output_dir = output_dir
        
        # Ensure output directory exists
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            self.logger.info(f"Created output directory: {output_dir}")
        
        # Check if openpyxl is available
        if not OPENPYXL_AVAILABLE:
            self.logger.warning("openpyxl not available, Excel generation will be limited")
        
        self.logger.info("DocumentGenerator initialized")
    
    def create_monitoring_excel(self, filename: str, content_type: str, data: List[Dict[str, Any]]) -> bool:
        """
        Create Excel file for monitoring data
        
        Args:
            filename: Name of the Excel file to create
            content_type: Type of content ('carousel', 'cancellation', 'news', 'media')
            data: List of data dictionaries to write to Excel
            
        Returns:
            bool: True if Excel file created successfully
        """
        try:
            if not OPENPYXL_AVAILABLE:
                self.logger.error("openpyxl not available, cannot create Excel file")
                return False
            
            if not data:
                self.logger.warning(f"No data provided for {content_type} Excel file")
                return False
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Create workbook and worksheet
            workbook = openpyxl.Workbook()
            worksheet = workbook.active
            worksheet.title = self._get_sheet_title(content_type)
            
            # Write headers
            headers = list(data[0].keys())
            self._write_headers(worksheet, headers)
            
            # Write data rows
            for row_idx, item in enumerate(data, start=2):
                for col_idx, header in enumerate(headers, start=1):
                    value = item.get(header, '')
                    # Convert datetime objects to strings
                    if isinstance(value, datetime):
                        value = value.strftime('%Y-%m-%d %H:%M:%S')
                    worksheet.cell(row=row_idx, column=col_idx, value=value)
            
            # Apply formatting
            self._apply_formatting(worksheet, len(headers), len(data))
            
            # Save workbook
            workbook.save(file_path)
            
            self.logger.info(f"Excel file created: {file_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error creating Excel file: {e}")
            return False
    
    def create_multi_sheet_excel(self, filename: str, sheets_data: Dict[str, List[Dict[str, Any]]]) -> bool:
        """
        Create Excel file with multiple sheets for different content types
        
        Args:
            filename: Name of the Excel file to create
            sheets_data: Dictionary mapping sheet names to data lists
            
        Returns:
            bool: True if multi-sheet Excel file created successfully
        """
        try:
            if not OPENPYXL_AVAILABLE:
                self.logger.error("openpyxl not available, cannot create Excel file")
                return False
            
            if not sheets_data:
                self.logger.warning("No sheet data provided for multi-sheet Excel file")
                return False
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Create workbook
            workbook = openpyxl.Workbook()
            
            # Remove default sheet
            workbook.remove(workbook.active)
            
            # Create sheets for each content type
            for sheet_name, data in sheets_data.items():
                if not data:
                    continue
                
                # Create worksheet
                worksheet = workbook.create_sheet(title=sheet_name)
                
                # Write headers
                headers = list(data[0].keys())
                self._write_headers(worksheet, headers)
                
                # Write data rows
                for row_idx, item in enumerate(data, start=2):
                    for col_idx, header in enumerate(headers, start=1):
                        value = item.get(header, '')
                        # Convert datetime objects to strings
                        if isinstance(value, datetime):
                            value = value.strftime('%Y-%m-%d %H:%M:%S')
                        worksheet.cell(row=row_idx, column=col_idx, value=value)
                
                # Apply formatting
                self._apply_formatting(worksheet, len(headers), len(data))
            
            # Save workbook
            workbook.save(file_path)
            
            self.logger.info(f"Multi-sheet Excel file created: {file_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error creating multi-sheet Excel file: {e}")
            return False
    
    def _get_sheet_title(self, content_type: str) -> str:
        """
        Get appropriate sheet title for content type
        
        Args:
            content_type: Type of content
            
        Returns:
            str: Sheet title
        """
        title_mapping = {
            'carousel': '輪播橫幅',
            'cancellation': '課程取消',
            'news': '新聞公告',
            'media': '多媒體內容'
        }
        
        return title_mapping.get(content_type, content_type.title())
    
    def _write_headers(self, worksheet, headers: List[str]):
        """
        Write headers to worksheet with formatting
        
        Args:
            worksheet: openpyxl worksheet
            headers: List of header strings
        """
        try:
            for col_idx, header in enumerate(headers, start=1):
                cell = worksheet.cell(row=1, column=col_idx, value=header)
                
                # Apply header formatting
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                
        except Exception as e:
            self.logger.warning(f"Error formatting headers: {e}")
    
    def _apply_formatting(self, worksheet, num_columns: int, num_rows: int):
        """
        Apply formatting to worksheet
        
        Args:
            worksheet: openpyxl worksheet
            num_columns: Number of columns
            num_rows: Number of data rows
        """
        try:
            # Auto-adjust column widths
            for col_idx in range(1, num_columns + 1):
                column_letter = openpyxl.utils.get_column_letter(col_idx)
                worksheet.column_dimensions[column_letter].width = 15
                
                # Special width adjustments for specific columns
                header_value = worksheet.cell(row=1, column=col_idx).value
                if header_value:
                    if 'URL' in str(header_value) or 'url' in str(header_value):
                        worksheet.column_dimensions[column_letter].width = 30
                    elif '內容' in str(header_value) or '描述' in str(header_value):
                        worksheet.column_dimensions[column_letter].width = 40
                    elif 'ID' in str(header_value) or 'id' in str(header_value):
                        worksheet.column_dimensions[column_letter].width = 20
            
            # Apply alternating row colors for better readability
            for row_idx in range(2, num_rows + 2):
                if row_idx % 2 == 0:
                    fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    for col_idx in range(1, num_columns + 1):
                        worksheet.cell(row=row_idx, column=col_idx).fill = fill
            
            # Apply alignment to all cells
            for row_idx in range(1, num_rows + 2):
                for col_idx in range(1, num_columns + 1):
                    cell = worksheet.cell(row=row_idx, column=col_idx)
                    cell.alignment = Alignment(vertical="center", wrap_text=True)
                    
        except Exception as e:
            self.logger.warning(f"Error applying formatting: {e}")
    
    def create_summary_excel(self, filename: str, summary_data: Dict[str, Any]) -> bool:
        """
        Create summary Excel file with monitoring statistics
        
        Args:
            filename: Name of the Excel file to create
            summary_data: Summary data dictionary
            
        Returns:
            bool: True if summary Excel file created successfully
        """
        try:
            if not OPENPYXL_AVAILABLE:
                self.logger.error("openpyxl not available, cannot create Excel file")
                return False
            
            file_path = os.path.join(self.output_dir, filename)
            
            # Create workbook and worksheet
            workbook = openpyxl.Workbook()
            worksheet = workbook.active
            worksheet.title = "監控摘要"
            
            # Write summary information
            row = 1
            
            # Title
            worksheet.cell(row=row, column=1, value="網站監控摘要報告")
            worksheet.cell(row=row, column=1).font = Font(size=16, bold=True)
            row += 2
            
            # Timestamp
            timestamp = summary_data.get('timestamp', datetime.now())
            worksheet.cell(row=row, column=1, value="生成時間:")
            worksheet.cell(row=row, column=2, value=timestamp.strftime('%Y-%m-%d %H:%M:%S'))
            row += 2
            
            # Content type statistics
            content_counts = summary_data.get('content_type_counts', {})
            if content_counts:
                worksheet.cell(row=row, column=1, value="內容類型統計:")
                worksheet.cell(row=row, column=1).font = Font(bold=True)
                row += 1
                
                for content_type, count in content_counts.items():
                    type_name = self._get_sheet_title(content_type)
                    worksheet.cell(row=row, column=2, value=f"{type_name}:")
                    worksheet.cell(row=row, column=3, value=count)
                    row += 1
                row += 1
            
            # Processing results
            processing_success = summary_data.get('processing_success', {})
            if processing_success:
                worksheet.cell(row=row, column=1, value="處理結果:")
                worksheet.cell(row=row, column=1).font = Font(bold=True)
                row += 1
                
                for content_type, success in processing_success.items():
                    type_name = self._get_sheet_title(content_type)
                    status = "成功" if success else "失敗"
                    worksheet.cell(row=row, column=2, value=f"{type_name}:")
                    worksheet.cell(row=row, column=3, value=status)
                    row += 1
            
            # Apply basic formatting
            for col in ['A', 'B', 'C']:
                worksheet.column_dimensions[col].width = 20
            
            # Save workbook
            workbook.save(file_path)
            
            self.logger.info(f"Summary Excel file created: {file_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error creating summary Excel file: {e}")
            return False
    
    def generate_both_documents(self, books_data: List[Dict[str, Any]], output_dir: str = None) -> tuple:
        """
        Generate both Word and Excel documents for book summaries
        
        Args:
            books_data: List of book dictionaries with summary information
            output_dir: Output directory (uses self.output_dir if None)
            
        Returns:
            tuple: (word_path, excel_path) paths to generated documents
        """
        try:
            if not books_data:
                raise ValueError("No books data provided")
            
            # Use provided output_dir or default
            target_dir = output_dir or self.output_dir
            
            # Ensure output directory exists
            if not os.path.exists(target_dir):
                os.makedirs(target_dir, exist_ok=True)
            
            # Generate timestamp for filenames
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            # Generate Word document
            word_filename = f"book_summaries_{timestamp}.docx"
            word_path = os.path.join(target_dir, word_filename)
            self._generate_word_document(books_data, word_path)
            
            # Generate Excel document
            excel_filename = f"book_summaries_{timestamp}.xlsx"
            excel_path = os.path.join(target_dir, excel_filename)
            self._generate_excel_document(books_data, excel_path)
            
            self.logger.info(f"Generated Word document: {word_path}")
            self.logger.info(f"Generated Excel document: {excel_path}")
            
            return word_path, excel_path
            
        except Exception as e:
            self.logger.error(f"Error generating documents: {e}")
            raise
    
    def _generate_word_document(self, books_data: List[Dict[str, Any]], output_path: str):
        """
        Generate Word document with book summaries
        
        Args:
            books_data: List of book dictionaries
            output_path: Path to save Word document
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Add title
            title = doc.add_heading('新書摘要報告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add generation date
            date_para = doc.add_paragraph()
            date_para.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Empty line
            
            # Add each book summary
            for idx, book in enumerate(books_data, 1):
                # Book title
                heading = doc.add_heading(f'{idx}. {book.get("title", "未知書名")}', level=1)
                
                # Book metadata
                meta_para = doc.add_paragraph()
                if book.get('author'):
                    meta_para.add_run(f'作者：{book["author"]}\n')
                meta_para.add_run(f'處理方法：{book.get("processing_method", "未知")}\n')
                meta_para.add_run(f'處理時間：{book.get("timestamp", "未知")}')
                
                # Summary
                doc.add_heading('摘要', level=2)
                summary_para = doc.add_paragraph(book.get('summary', '無摘要'))
                
                # Add separator
                if idx < len(books_data):
                    doc.add_paragraph('─' * 50)
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"Word document saved: {output_path}")
            
        except ImportError:
            self.logger.error("python-docx not available, cannot create Word document")
            raise
        except Exception as e:
            self.logger.error(f"Error creating Word document: {e}")
            raise
    
    def _generate_excel_document(self, books_data: List[Dict[str, Any]], output_path: str):
        """
        Generate Excel document with book summaries
        
        Args:
            books_data: List of book dictionaries
            output_path: Path to save Excel document
        """
        try:
            if not OPENPYXL_AVAILABLE:
                self.logger.error("openpyxl not available, cannot create Excel document")
                raise ImportError("openpyxl not available")
            
            # Create workbook
            workbook = openpyxl.Workbook()
            worksheet = workbook.active
            worksheet.title = "新書摘要"
            
            # Define headers
            headers = ['序號', '書名', '作者', '摘要', '處理方法', '處理時間', 'PDF URL']
            
            # Write headers
            self._write_headers(worksheet, headers)
            
            # Write data
            for idx, book in enumerate(books_data, 1):
                row_data = [
                    idx,
                    book.get('title', ''),
                    book.get('author', ''),
                    book.get('summary', ''),
                    book.get('processing_method', ''),
                    book.get('timestamp', ''),
                    book.get('pdf_url', '')
                ]
                
                for col_idx, value in enumerate(row_data, 1):
                    worksheet.cell(row=idx+1, column=col_idx, value=value)
            
            # Apply formatting
            self._apply_formatting(worksheet, len(headers), len(books_data))
            
            # Adjust column widths for book summaries
            worksheet.column_dimensions['B'].width = 30  # Title
            worksheet.column_dimensions['C'].width = 20  # Author
            worksheet.column_dimensions['D'].width = 50  # Summary
            worksheet.column_dimensions['G'].width = 40  # URL
            
            # Save workbook
            workbook.save(output_path)
            self.logger.info(f"Excel document saved: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Error creating Excel document: {e}")
            raise
    
    def get_output_directory(self) -> str:
        """
        Get the output directory path
        
        Returns:
            str: Output directory path
        """
        return self.output_dir
    
    def list_generated_files(self) -> List[str]:
        """
        List all generated files in the output directory
        
        Returns:
            List[str]: List of generated file paths
        """
        try:
            if not os.path.exists(self.output_dir):
                return []
            
            files = []
            for filename in os.listdir(self.output_dir):
                if filename.endswith(('.xlsx', '.xls')):
                    file_path = os.path.join(self.output_dir, filename)
                    files.append(file_path)
            
            return sorted(files, key=os.path.getmtime, reverse=True)
            
        except Exception as e:
            self.logger.error(f"Error listing generated files: {e}")
            return []
    
    def cleanup_old_files(self, days_old: int = 30) -> int:
        """
        Clean up old generated files
        
        Args:
            days_old: Remove files older than this many days
            
        Returns:
            int: Number of files removed
        """
        try:
            if not os.path.exists(self.output_dir):
                return 0
            
            import time
            current_time = time.time()
            cutoff_time = current_time - (days_old * 24 * 60 * 60)
            
            removed_count = 0
            
            for filename in os.listdir(self.output_dir):
                if filename.endswith(('.xlsx', '.xls')):
                    file_path = os.path.join(self.output_dir, filename)
                    
                    try:
                        file_time = os.path.getmtime(file_path)
                        if file_time < cutoff_time:
                            os.remove(file_path)
                            removed_count += 1
                            self.logger.info(f"Removed old file: {filename}")
                    except Exception as e:
                        self.logger.warning(f"Error removing file {filename}: {e}")
            
            if removed_count > 0:
                self.logger.info(f"Cleaned up {removed_count} old files")
            
            return removed_count
            
        except Exception as e:
            self.logger.error(f"Error cleaning up old files: {e}")
            return 0


# Example usage and testing
def main():
    """
    Example usage of DocumentGenerator
    """
    import logging
    
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    logger = logging.getLogger(__name__)
    
    try:
        # Initialize DocumentGenerator
        generator = DocumentGenerator(logger=logger)
        
        # Test data
        test_data = [
            {
                'ID': 'test_1',
                '標題': '測試項目 1',
                '內容': '這是測試內容',
                '時間': datetime.now()
            },
            {
                'ID': 'test_2',
                '標題': '測試項目 2',
                '內容': '這是另一個測試內容',
                '時間': datetime.now()
            }
        ]
        
        # Test single sheet Excel creation
        success = generator.create_monitoring_excel(
            filename="test_monitoring.xlsx",
            content_type="news",
            data=test_data
        )
        logger.info(f"Single sheet Excel creation: {success}")
        
        # Test multi-sheet Excel creation
        sheets_data = {
            '新聞公告': test_data,
            '輪播橫幅': test_data
        }
        
        success = generator.create_multi_sheet_excel(
            filename="test_multi_sheet.xlsx",
            sheets_data=sheets_data
        )
        logger.info(f"Multi-sheet Excel creation: {success}")
        
        # Test summary Excel creation
        summary_data = {
            'timestamp': datetime.now(),
            'content_type_counts': {'news': 2, 'carousel': 1},
            'processing_success': {'news': True, 'carousel': True}
        }
        
        success = generator.create_summary_excel(
            filename="test_summary.xlsx",
            summary_data=summary_data
        )
        logger.info(f"Summary Excel creation: {success}")
        
        # List generated files
        files = generator.list_generated_files()
        logger.info(f"Generated files: {files}")
        
    except Exception as e:
        logger.error(f"Error during testing: {e}")


if __name__ == "__main__":
    main()